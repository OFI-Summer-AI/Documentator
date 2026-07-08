from __future__ import annotations

import base64
import json
import io
import os
import re
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from pypdf import PdfReader

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from django.utils.text import slugify
from openai import OpenAI
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    HRFlowable, Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
)
from reportlab.platypus.tableofcontents import TableOfContents

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

ASSETS_DIR = Path(__file__).resolve().parents[2] / "assets"

DocSection = dict[str, Any]

# ---------------------------------------------------------------------------
# Shared design tokens — one palette/scale reused across the PDF, DOCX, and
# LaTeX renderers so the three output formats read as the same document.
# ---------------------------------------------------------------------------

PALETTE = {
    "ink": "0f3d4a",  # cover title / darkest text
    "ink_text": "1f2933",  # body copy
    "accent": "0e7490",  # primary teal — headings, table headers, rules
    "accent_dark": "0b4f5c",  # cover band background
    "gold": "d7aa3a",  # secondary accent — thin rules, callout border
    "muted": "415d66",  # subtitles, captions
    "muted_light": "5b7280",  # footer text
    "border": "c9e4e8",  # table/box borders
    "border_light": "dbeaec",  # table inner gridlines
    "band": "f5fafb",  # table banding / callout background
    "white": "ffffff",
}


def _hex(name: str) -> str:
    """Palette colour as a reportlab-style '#RRGGBB' string."""
    return f"#{PALETTE[name]}"


def _numbered_titles(document_sections: list[DocSection]) -> dict[int, str]:
    """Map section index -> 'N. Title' for every titled (non-figure) section, in document order."""
    numbering: dict[int, str] = {}
    counter = 0
    for i, sec in enumerate(document_sections):
        title = str(sec.get("title", ""))
        if not title or sec.get("type") == "figure":
            continue
        counter += 1
        numbering[i] = f"{counter}. {title}"
    return numbering

_LENGTH_GUIDANCE = {
    "brief": (
        "Keep the document concise: aim for 3-5 sections total (including the executive summary and version "
        "control table), short paragraphs (2-4 sentences), and short bullet lists (3-5 items)."
    ),
    "standard": (
        "Aim for a standard business-document length: 5-8 sections total, paragraphs of moderate depth "
        "(4-8 sentences), and bullet lists sized to the content."
    ),
    "detailed": (
        "Produce a thorough, detailed document: 8-12 sections total, in-depth paragraphs (8+ sentences) with "
        "supporting detail, and comprehensive bullet/table breakdowns."
    ),
}

_TONE_GUIDANCE = {
    "formal": "Write in a formal, professional register suitable for executive stakeholders.",
    "consulting": (
        "Write in a confident management-consulting tone: crisp, structured, benefits-oriented, using "
        "industry-standard framing."
    ),
    "technical": (
        "Write in a precise technical tone aimed at engineers/architects: exact terminology, specific and "
        "unambiguous statements."
    ),
    "casual": "Write in a clear, approachable tone while still being professional; avoid unnecessary jargon.",
}

_TEMPLATE_GUIDANCE = {
    "general": "",
    "meeting_minutes": (
        "Structure this as formal meeting minutes. After the executive summary, include (in order): an "
        "Attendees section (bullets listing attendee names/roles mentioned, or 'Not specified' if none are "
        "found), a Discussion Summary, an Action Items section (bullets or a table with owner and due date "
        "when known), and a Decisions Made section if any decisions were mentioned."
    ),
    "technical_spec": (
        "Structure this as a technical specification. After the executive summary, include (in order): "
        "Overview/Background, Requirements (bullets or a table), Architecture/Design, and Risks & Open "
        "Questions. Use tables for structured requirements or comparisons."
    ),
    "proposal": (
        "Structure this as a client proposal. After the executive summary, include (in order): Scope of Work, "
        "Timeline (prefer a table with milestones/dates), Pricing/Investment (a table if figures are "
        "available, otherwise a short bullets/paragraph placeholder), and Next Steps."
    ),
    "status_report": (
        "Structure this as a project status report. After the executive summary, include (in order): Progress "
        "Since Last Update, Current Status (consider a status table), Blockers & Risks, and Next Steps."
    ),
}

_TEMPLATE_SECTION_LABELS = {
    "meeting_minutes": ["Attendees", "Discussion Summary", "Action Items"],
    "technical_spec": ["Overview", "Requirements", "Risks & Open Questions"],
    "proposal": ["Scope of Work", "Timeline", "Next Steps"],
    "status_report": ["Progress Since Last Update", "Blockers & Risks", "Next Steps"],
}


def _length_guidance(target_length: str) -> str:
    return _LENGTH_GUIDANCE.get(target_length, _LENGTH_GUIDANCE["standard"])


def _tone_guidance(tone: str) -> str:
    return _TONE_GUIDANCE.get(tone, _TONE_GUIDANCE["formal"])


def _template_guidance(template: str) -> str:
    return _TEMPLATE_GUIDANCE.get(template, "")


def _apply_max_sections(sections: list[DocSection], max_sections: int | None) -> list[DocSection]:
    """Trim a section list down to max_sections, always preserving figures and the final section
    (typically the version-control table)."""
    if not max_sections or len(sections) <= max_sections:
        return sections

    protected_ids = {id(s) for s in sections if s.get("type") == "figure"}
    if sections:
        protected_ids.add(id(sections[-1]))

    fillable = [s for s in sections if id(s) not in protected_ids]
    budget = max(max_sections - len(protected_ids), 0)
    keep_ids = {id(s) for s in fillable[:budget]} | protected_ids
    return [s for s in sections if id(s) in keep_ids]


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass(slots=True)
class RenderedDocument:
    title: str
    client_name: str
    source_text: str
    document_sections: list[DocSection]
    filename: str
    pdf_bytes: bytes
    docx_bytes: bytes
    tex_zip_bytes: bytes
    latex_source: str
    generation_mode: str


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def read_image_data(source_images: list[Any], image_descriptions: list[Any]) -> list[dict[str, Any]]:
    image_data: list[dict[str, Any]] = []
    for i, img_file in enumerate(source_images):
        desc = image_descriptions[i] if i < len(image_descriptions) else ""
        try:
            img_file.seek(0)
            image_data.append({"bytes": img_file.read(), "description": str(desc).strip()})
        except Exception:
            pass
    return image_data


def generate_sections(validated_data: dict[str, Any]) -> dict[str, Any]:
    """Run source ingestion + LLM/fallback generation and return the structured document data,
    without rendering PDF/DOCX/LaTeX. Used by both the one-shot flow and the review-before-export flow."""
    source_text = validated_data["source_text"].strip()
    agent_instructions = validated_data.get("agent_instructions", "").strip()
    document_language = detect_document_language(source_text, agent_instructions)

    pdf_texts = extract_pdf_texts(validated_data.get("source_pdfs") or [])
    if pdf_texts:
        joined = "\n\n---\n\n".join(pdf_texts)
        source_text = f"{source_text}\n\n--- Extracted from uploaded PDFs ---\n\n{joined}".strip()

    # Read uploaded images into memory before any async/generator usage
    source_images = validated_data.get("source_images") or []
    image_descriptions = validated_data.get("image_descriptions") or []
    image_data = read_image_data(source_images, image_descriptions)

    template = validated_data.get("template") or "general"
    target_length = validated_data.get("target_length") or "standard"
    tone = validated_data.get("tone") or "formal"
    max_sections = validated_data.get("max_sections")

    doc_data = generate_document_data(
        source_text,
        agent_instructions,
        document_language,
        image_data,
        template=template,
        target_length=target_length,
        tone=tone,
        max_sections=max_sections,
    )
    generation_mode = str(doc_data.get("generation_mode", "fallback"))
    title = str(doc_data.get("title", "")).strip() or _fallback_title(source_text)
    client_name = str(doc_data.get("client_name", "")).strip()
    document_sections: list[DocSection] = doc_data.get("document_sections", [])

    if not document_sections:
        document_sections = _fallback_sections(
            source_text, document_language, image_data, template=template, target_length=target_length,
            max_sections=max_sections,
        )
    elif image_data:
        document_sections = _ensure_figures(document_sections, image_data)

    return {
        "title": title,
        "client_name": client_name,
        "source_text": source_text,
        "document_language": document_language,
        "document_sections": document_sections,
        "generation_mode": generation_mode,
        "image_data": image_data,
        "filename": slugify(title) or "documentation",
    }


def render_sections(
    *,
    title: str,
    client_name: str,
    source_text: str,
    document_sections: list[DocSection],
    document_language: str,
    logo: Any,
    image_data: list[dict[str, Any]] | None = None,
    generation_mode: str = "fallback",
    filename: str | None = None,
) -> RenderedDocument:
    """Render PDF/DOCX/LaTeX from already-generated (and possibly user-edited) document sections.

    The LaTeX/PDF/DOCX layouts are always built deterministically from document_sections (never from
    LLM-authored LaTeX) so every export uses the same polished, consistent template."""
    image_data = image_data or []
    latex_source = _build_latex(title, client_name, document_sections, document_language, image_data)
    filename = filename or slugify(title) or "documentation"
    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    pdf_bytes = render_pdf(
        title=title,
        client_name=client_name,
        source_text=source_text,
        document_sections=document_sections,
        document_language=document_language,
        logo=logo,
        timestamp=timestamp,
        image_data=image_data,
    )
    docx_bytes = render_docx(
        title=title,
        client_name=client_name,
        document_sections=document_sections,
        document_language=document_language,
        timestamp=timestamp,
        image_data=image_data,
    )
    tex_zip_bytes = build_tex_zip(
        latex_source=latex_source,
        filename=filename,
        image_data=image_data,
    )

    return RenderedDocument(
        title=title,
        client_name=client_name,
        source_text=source_text,
        document_sections=document_sections,
        filename=filename,
        pdf_bytes=pdf_bytes,
        docx_bytes=docx_bytes,
        tex_zip_bytes=tex_zip_bytes,
        latex_source=latex_source,
        generation_mode=generation_mode,
    )


def build_document_payload(validated_data: dict[str, Any]) -> RenderedDocument:
    gen = generate_sections(validated_data)
    return render_sections(
        title=gen["title"],
        client_name=gen["client_name"],
        source_text=gen["source_text"],
        document_sections=gen["document_sections"],
        document_language=gen["document_language"],
        logo=validated_data.get("logo"),
        image_data=gen["image_data"],
        generation_mode=gen["generation_mode"],
        filename=gen["filename"],
    )


# ---------------------------------------------------------------------------
# LLM generation
# ---------------------------------------------------------------------------

_SYSTEM_PROMPT = (
    "You are an expert technical writer. Transform the user-provided transcript or notes into a "
    "polished, client-ready document. You have COMPLETE freedom to decide:\n"
    "- How many sections the document needs.\n"
    "- The title and content of every section.\n"
    "- Whether each section is best represented as a paragraph of text, a bullet list, or a table.\n\n"
    "Rules:\n"
    "1. Never paste raw transcript text into the document. Always synthesise and rewrite.\n"
    "2. Use tables when comparing options, listing attributes, or showing structured data.\n"
    "3. Use bullet lists for action items, requirements, or enumerable facts.\n"
    "4. Use paragraphs for narrative explanations, summaries, and analysis.\n"
    "5. Always include a short executive summary as the first section.\n"
    "6. Always include a version/control table as the last section.\n"
    "7. If the transcript mentions a client name, extract it.\n\n"
    "Return ONLY a valid JSON object with this exact schema (no markdown fences):\n"
    '{\n'
    '  "title": "<concise document title>",\n'
    '  "client_name": "<client name or empty string>",\n'
    '  "document_sections": [\n'
    '    {"title": "<section title>", "type": "paragraph", "content": "<synthesised prose>"},\n'
    '    {"title": "<section title>", "type": "bullets", "content": ["<item>", "<item>"]},\n'
    '    {"title": "<section title>", "type": "table", "content": {"headers": ["<col1>", "<col2>"], "rows": [["<val>", "<val>"]]}},\n'
    '    {"title": "", "type": "figure", "content": {"figure_index": 0, "figure_number": 1, "caption": "<Figure N: descriptive caption>"}}\n'
    '  ]\n'
    '}'
)


def generate_document_data(
    source_text: str,
    agent_instructions: str = "",
    document_language: str = "en",
    image_data: list[dict[str, Any]] | None = None,
    *,
    template: str = "general",
    target_length: str = "standard",
    tone: str = "formal",
    max_sections: int | None = None,
) -> dict[str, Any]:
    api_key = _env_value("OPENAI_API_KEY")
    model = _env_value("OPENAI_MODEL", default="gpt-4.1-mini")

    if not api_key:
        return _fallback_document_data(
            source_text, document_language, image_data or [], template=template, target_length=target_length,
            max_sections=max_sections,
        )

    try:
        client = OpenAI(api_key=api_key)
        messages: list[dict[str, str]] = [{"role": "system", "content": _SYSTEM_PROMPT}]
        messages.append({
            "role": "system",
            "content": (
                "Write the section titles and content in Spanish if the transcript/instructions are mainly Spanish. "
                "Write them in English if the transcript/instructions are mainly English. "
                f"Detected language preference: {document_language}."
            ),
        })
        guidance_parts = [_length_guidance(target_length), _tone_guidance(tone), _template_guidance(template)]
        if max_sections:
            guidance_parts.append(f"Do not exceed {max_sections} sections in total.")
        messages.append({
            "role": "system",
            "content": " ".join(part for part in guidance_parts if part),
        })
        if image_data:
            img_list = "\n".join(
                f"  - figure_index={i}: {d['description'] or '(no description)'}"
                for i, d in enumerate(image_data)
            )
            messages.append({
                "role": "system",
                "content": (
                    f"MANDATORY: The user has uploaded {len(image_data)} image(s) that MUST appear in the "
                    f"document_sections array. Every figure_index from 0 to {len(image_data) - 1} must have "
                    "exactly one corresponding section of type 'figure' in the output.\n\n"
                    f"Available images:\n{img_list}\n\n"
                    "For each image, decide the best position in the document (where the image most supports "
                    "the surrounding content) and insert a figure section there. Write a professional caption "
                    "that describes what the image shows — it does not need to copy the user's description verbatim.\n\n"
                    "Figure section format (title must be empty string):\n"
                    '{"title": "", "type": "figure", "content": {"figure_index": 0, "figure_number": 1, "caption": "Figure 1: Architecture overview"}}\n\n'
                    "Number figures sequentially starting at 1. Omitting any image is NOT allowed."
                ),
            })
        if agent_instructions:
            messages.append({"role": "system", "content": f"Additional author instructions: {agent_instructions}"})
        messages.append({"role": "user", "content": source_text})

        response = client.chat.completions.create(
            model=model,
            temperature=0.3,
            response_format={"type": "json_object"},
            messages=messages,
        )
        raw = response.choices[0].message.content or "{}"
        parsed = json.loads(raw)
        parsed["generation_mode"] = "openai"
        sections = _normalise_sections(
            parsed.get("document_sections", []), source_text, document_language, image_data or []
        )
        parsed["document_sections"] = _apply_max_sections(sections, max_sections)
        return parsed
    except Exception:
        return _fallback_document_data(
            source_text, document_language, image_data or [], template=template, target_length=target_length,
            max_sections=max_sections,
        )


def regenerate_section(
    *,
    source_text: str,
    agent_instructions: str = "",
    document_language: str = "en",
    template: str = "general",
    target_length: str = "standard",
    tone: str = "formal",
    section_title: str,
    section_type: str = "paragraph",
    section_instructions: str = "",
) -> DocSection:
    """Regenerate a single section's content, used by the review-before-export flow."""
    api_key = _env_value("OPENAI_API_KEY")
    model = _env_value("OPENAI_MODEL", default="gpt-4.1-mini")

    if not api_key:
        return _fallback_regenerate_section(source_text, document_language, section_title, section_type)

    try:
        client = OpenAI(api_key=api_key)
        guidance = " ".join(part for part in (
            _length_guidance(target_length), _tone_guidance(tone), _template_guidance(template),
        ) if part)
        system_prompt = (
            "You are revising a single section of a client-ready document. Return ONLY a valid JSON object "
            "(no markdown fences) for exactly one section using this schema:\n"
            '{"title": "<title>", "type": "paragraph|bullets|table", '
            '"content": "<prose>" | ["<item>", "<item>"] | {"headers": ["<col>"], "rows": [["<val>"]]}}\n\n'
            f"Section to write: \"{section_title}\". The \"type\" MUST be \"{section_type}\" — keep the same "
            f"presentation format as before, matching \"content\" to that type's shape. Never paste raw "
            f"transcript text — always synthesise and rewrite. {guidance}"
        )
        messages: list[dict[str, str]] = [
            {"role": "system", "content": system_prompt},
            {
                "role": "system",
                "content": f"Write the section in {'Spanish' if document_language == 'es' else 'English'}.",
            },
        ]
        if agent_instructions:
            messages.append({"role": "system", "content": f"Additional author instructions: {agent_instructions}"})
        if section_instructions:
            messages.append({"role": "system", "content": f"Specific instructions for this section: {section_instructions}"})
        messages.append({"role": "user", "content": source_text})

        response = client.chat.completions.create(
            model=model,
            temperature=0.4,
            response_format={"type": "json_object"},
            messages=messages,
        )
        raw = response.choices[0].message.content or "{}"
        parsed = json.loads(raw)
        normalised = _normalise_sections([parsed], source_text, document_language, [])
        if normalised:
            section = normalised[0]
            section["title"] = section.get("title") or section_title
            return _coerce_section_type(section, section_type)
        return _fallback_regenerate_section(source_text, document_language, section_title, section_type)
    except Exception:
        return _fallback_regenerate_section(source_text, document_language, section_title, section_type)


def _coerce_section_type(section: DocSection, target_type: str) -> DocSection:
    """Force a regenerated section back to the type the caller asked for, converting content shape
    if the LLM drifted (e.g. returned a paragraph when bullets were requested)."""
    if section.get("type") == target_type or target_type not in ("paragraph", "bullets", "table"):
        return section

    content = section.get("content")
    if target_type == "bullets":
        if isinstance(content, list):
            new_content: Any = content
        elif isinstance(content, dict):
            new_content = [" — ".join(row) for row in content.get("rows", [])] or [str(content)]
        else:
            parts = re.split(r"(?<=[.!?])\s+", str(content).strip())
            new_content = [p for p in parts if p] or [str(content)]
    elif target_type == "paragraph":
        if isinstance(content, list):
            new_content = " ".join(str(item) for item in content)
        elif isinstance(content, dict):
            rows = content.get("rows", [])
            new_content = " ".join(" — ".join(row) for row in rows)
        else:
            new_content = str(content)
    else:  # table
        if isinstance(content, dict):
            new_content = content
        elif isinstance(content, list):
            new_content = {"headers": ["Item"], "rows": [[str(item)] for item in content]}
        else:
            new_content = {"headers": ["Content"], "rows": [[str(content)]]}

    return {"title": section.get("title", ""), "type": target_type, "content": new_content}


def _fallback_regenerate_section(
    source_text: str,
    document_language: str,
    section_title: str,
    section_type: str,
) -> DocSection:
    if section_type == "bullets":
        content: Any = _extract_lines(source_text)[:6] or ["See source context for details."]
    elif section_type == "table":
        labels = localized_labels(document_language)
        content = {
            "headers": labels["version_headers"],
            "rows": [["1.0", datetime.now(timezone.utc).strftime("%Y-%m-%d"), "", labels["initial_draft"]]],
        }
    else:
        content = _fallback_summary(source_text, sentence_count=2)
    return {"title": section_title, "type": section_type, "content": content}


# ---------------------------------------------------------------------------
# Section normalisation & fallback
# ---------------------------------------------------------------------------

def _ensure_figures(sections: list[DocSection], image_data: list[dict[str, Any]]) -> list[DocSection]:
    """Guarantee every uploaded image has a figure section in the document.
    Any image the LLM forgot to place is injected before the last section."""
    placed = {
        int(s["content"]["figure_index"])
        for s in sections
        if s.get("type") == "figure" and isinstance(s.get("content"), dict)
    }
    missing = [i for i in range(len(image_data)) if i not in placed]
    if not missing:
        return sections

    # Count already-placed figures to continue numbering correctly
    next_fig_num = sum(1 for s in sections if s.get("type") == "figure") + 1
    inject = []
    for i in missing:
        img = image_data[i]
        caption = f"Figure {next_fig_num}: {img['description']}" if img["description"] else f"Figure {next_fig_num}"
        inject.append({
            "title": "",
            "type": "figure",
            "content": {"figure_index": i, "figure_number": next_fig_num, "caption": caption},
        })
        next_fig_num += 1

    # Insert before the last section (typically the version control table)
    insert_pos = max(len(sections) - 1, 0)
    return sections[:insert_pos] + inject + sections[insert_pos:]


def _normalise_sections(
    raw: Any,
    source_text: str,
    document_language: str,
    image_data: list[dict[str, Any]],
) -> list[DocSection]:
    if not isinstance(raw, list):
        return _fallback_sections(source_text, document_language, image_data)

    sections: list[DocSection] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        sec_type = str(item.get("type", "paragraph")).strip().lower()
        title = str(item.get("title", "")).strip()
        content = item.get("content", "")

        if sec_type == "bullets":
            content = [str(b).strip() for b in (content if isinstance(content, list) else [str(content)]) if str(b).strip()]
            if not content:
                continue
        elif sec_type == "table":
            if not isinstance(content, dict):
                continue
            headers = [str(h) for h in content.get("headers", [])]
            rows = [[str(c) for c in row] for row in content.get("rows", [])]
            if not headers:
                continue
            content = {"headers": headers, "rows": rows}
        elif sec_type == "figure":
            if not isinstance(content, dict):
                continue
            try:
                figure_index = int(content.get("figure_index", 0))
            except (TypeError, ValueError):
                figure_index = 0
            try:
                figure_number = int(content.get("figure_number", 1))
            except (TypeError, ValueError):
                figure_number = 1
            caption = str(content.get("caption", f"Figure {figure_number}")).strip()
            if figure_index >= len(image_data):
                continue
            content = {"figure_index": figure_index, "figure_number": figure_number, "caption": caption}
        else:
            sec_type = "paragraph"
            content = str(content).strip()
            if not content:
                continue

        sections.append({"title": title, "type": sec_type, "content": content})

    return sections if sections else _fallback_sections(source_text, document_language, image_data)


_LENGTH_KEY_POINT_COUNT = {"brief": 4, "standard": 6, "detailed": 10}
_LENGTH_SUMMARY_SENTENCES = {"brief": 1, "standard": 2, "detailed": 4}


def _fallback_sections(
    source_text: str,
    document_language: str,
    image_data: list[dict[str, Any]] | None = None,
    *,
    template: str = "general",
    target_length: str = "standard",
    max_sections: int | None = None,
) -> list[DocSection]:
    summary = _fallback_summary(source_text, sentence_count=_LENGTH_SUMMARY_SENTENCES.get(target_length, 2))
    point_count = _LENGTH_KEY_POINT_COUNT.get(target_length, 6)
    lines = _extract_lines(source_text)[:point_count] or ["See source context for details."]
    labels = localized_labels(document_language)
    sections: list[DocSection] = [
        {"title": labels["executive_summary"], "type": "paragraph", "content": summary},
        {"title": labels["key_points"], "type": "bullets", "content": lines},
    ]
    sections.extend(_template_fallback_sections(template, source_text, document_language))
    for i, img in enumerate(image_data or []):
        caption = f"Figure {i + 1}: {img['description']}" if img["description"] else f"Figure {i + 1}"
        sections.append({
            "title": "",
            "type": "figure",
            "content": {"figure_index": i, "figure_number": i + 1, "caption": caption},
        })
    sections.append({
        "title": labels["version_control"],
        "type": "table",
        "content": {
            "headers": labels["version_headers"],
            "rows": [["1.0", datetime.now(timezone.utc).strftime("%Y-%m-%d"), "", labels["initial_draft"]]],
        },
    })
    return _apply_max_sections(sections, max_sections)


def _template_fallback_sections(template: str, source_text: str, document_language: str) -> list[DocSection]:
    """Deterministic skeleton sections added per template when generating without an LLM."""
    labels = _TEMPLATE_SECTION_LABELS.get(template)
    if not labels:
        return []
    lines = _extract_lines(source_text)
    placeholder = "Pendiente de detalle." if document_language == "es" else "To be detailed."
    sections: list[DocSection] = []
    for label in labels:
        sections.append({"title": label, "type": "bullets", "content": lines[:3] or [placeholder]})
    return sections


def _fallback_document_data(
    source_text: str,
    document_language: str,
    image_data: list[dict[str, Any]] | None = None,
    *,
    template: str = "general",
    target_length: str = "standard",
    max_sections: int | None = None,
) -> dict[str, Any]:
    title = _fallback_title(source_text)
    client_name = _fallback_client_name(source_text)
    document_sections = _fallback_sections(
        source_text, document_language, image_data, template=template, target_length=target_length,
        max_sections=max_sections,
    )
    return {
        "title": title,
        "client_name": client_name,
        "document_sections": document_sections,
        "generation_mode": "fallback",
    }


# ---------------------------------------------------------------------------
# LaTeX generation
# ---------------------------------------------------------------------------

def _build_latex(
    title: str,
    client_name: str,
    document_sections: list[DocSection],
    document_language: str,
    image_data: list[dict[str, Any]] | None = None,
) -> str:
    logo_path = resolve_ofi_logo_path()
    labels = localized_labels(document_language)
    ofi_logo_block = "\\includegraphics[height=1cm]{ofi-logo}" if logo_path else "\\textbf{\\textcolor{docaccent}{OFI}}"
    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    numbering = _numbered_titles(document_sections)
    has_figures = any(s.get("type") == "figure" for s in document_sections)

    def heading(index: int, use_box: bool = False) -> str:
        """A numbered, TOC-registered heading. `use_box` renders it as the lead summary callout."""
        numbered_title = numbering.get(index, "")
        if use_box:
            return (
                f"\\phantomsection\\addcontentsline{{toc}}{{section}}{{{_le(numbered_title)}}}\n"
            )
        return (
            f"\\phantomsection\\addcontentsline{{toc}}{{section}}{{{_le(numbered_title)}}}\n"
            f"\\section*{{{_le(numbered_title)}}}\n"
        )

    body_blocks: list[str] = []
    for i, sec in enumerate(document_sections):
        sec_title = str(sec.get("title", ""))
        sec_type = str(sec.get("type", "paragraph"))
        content = sec.get("content", "")
        is_lead_summary = i == 0 and sec_type == "paragraph" and sec_title

        if sec_type == "figure" and isinstance(content, dict):
            caption = str(content.get("caption", ""))
            fig_idx = content.get("figure_index", 0)
            body_blocks.append(
                "\\begin{figure}[H]\n"
                "\\centering\n"
                f"\\fcolorbox{{docborder}}{{white}}{{\\includegraphics[width=0.78\\textwidth]{{figure_{fig_idx}}}}}\\\\[0.25cm]\n"
                f"{{\\small\\color{{docmuted}}{_le(caption)}}}\n"
                "\\end{figure}"
            )
        elif is_lead_summary:
            body_blocks.append(
                heading(i, use_box=True)
                + "\\begin{tcolorbox}[colback=docband, colframe=docborder, coltitle=white, "
                "colbacktitle=docaccent, fonttitle=\\bfseries, title={" + _le(numbering.get(i, sec_title))
                + "}, boxrule=0.4pt, arc=2mm, left=10pt, right=10pt, top=8pt, bottom=8pt]\n"
                f"{_le(str(content))}\n"
                "\\end{tcolorbox}"
            )
        elif sec_type == "bullets":
            items_str = "\n".join(
                f"\\item {_le(str(b))}" for b in (content if isinstance(content, list) else [str(content)])
            )
            body_blocks.append(
                heading(i)
                + f"\\begin{{itemize}}[leftmargin=1.2em, itemsep=3pt, label=\\textcolor{{docaccent}}{{\\textbullet}}]\n"
                f"{items_str}\n\\end{{itemize}}"
            )
        elif sec_type == "table" and isinstance(content, dict):
            headers = content.get("headers", [])
            rows = content.get("rows", [])
            col_count = max(len(headers), max((len(r) for r in rows), default=0), 1)
            col_fmt = "".join(["X"] * col_count)
            header_row = " & ".join(f"\\textcolor{{white}}{{\\textbf{{{_le(str(h))}}}}}" for h in headers) + " \\\\"
            data_rows = "\n".join(" & ".join(_le(str(c)) for c in row) + " \\\\" for row in rows)
            body_blocks.append(
                heading(i)
                + "{\\renewcommand{\\arraystretch}{1.35}\\rowcolors{2}{docband}{white}\n"
                f"\\begin{{tabularx}}{{\\textwidth}}{{{col_fmt}}}\n"
                f"\\toprule\n\\rowcolor{{docaccent}} {header_row}\n\\midrule\n{data_rows}\n\\bottomrule\n"
                "\\end{tabularx}}"
            )
        else:
            if sec_title:
                body_blocks.append(heading(i) + _le(str(content)))
            else:
                body_blocks.append(_le(str(content)))

    body_block = "\n\n".join(body_blocks)
    float_pkg = "\\usepackage{float}\n" if has_figures else ""
    client_line = (
        f"{{\\color{{white}}\\large {_le(labels['client'])}: {_le(client_name)}}}\\\\[0.3cm]\n"
        if client_name
        else ""
    )
    header_title = title if len(title) <= 60 else f"{title[:57]}..."

    color_defs = "\n".join(
        f"\\definecolor{{doc{name.replace('_', '')}}}{{HTML}}{{{hexval.upper()}}}" for name, hexval in PALETTE.items()
    )

    return (
        "\\documentclass[11pt,a4paper]{article}\n"
        "\\usepackage[margin=2.2cm]{geometry}\n"
        "\\usepackage[T1]{fontenc}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage{lmodern}\n"
        "\\usepackage{parskip}\n"
        "\\usepackage{graphicx}\n"
        f"{float_pkg}"
        "\\usepackage{tabularx}\n"
        "\\usepackage{booktabs}\n"
        "\\usepackage{enumitem}\n"
        "\\usepackage[table]{xcolor}\n"
        f"{color_defs}\n"
        "\\usepackage{titlesec}\n"
        "\\usepackage{tcolorbox}\n"
        "\\usepackage{fancyhdr}\n"
        "\\usepackage{hyperref}\n"
        "\\hypersetup{colorlinks=true, linkcolor=docaccent, urlcolor=docaccent, "
        f"pdftitle={{{_le(title)}}}, pdfauthor={{Documentator}}}}\n"
        "\\titleformat{\\section}{\\Large\\bfseries\\color{docink}}{}{0pt}{}"
        "[{\\color{docgold}\\titlerule[1.1pt]}\\vspace{2pt}]\n"
        "\\titlespacing*{\\section}{0pt}{1.4em}{0.9em}\n"
        "\\pagestyle{fancy}\n"
        "\\fancyhf{}\n"
        f"\\fancyhead[L]{{\\small\\color{{docmutedlight}} {_le(header_title)}}}\n"
        f"\\fancyhead[R]{{{ofi_logo_block}}}\n"
        "\\fancyfoot[L]{\\small\\color{docmutedlight}\\textit{Documentator}}\n"
        "\\fancyfoot[R]{\\small\\color{docmutedlight}\\thepage}\n"
        "\\renewcommand{\\headrulewidth}{0.6pt}\n"
        "\\renewcommand{\\footrulewidth}{0.4pt}\n"
        "\\setlength{\\parindent}{0pt}\n"
        "\\begin{document}\n"
        "\\begin{titlepage}\n"
        "\\centering\n"
        "\\vspace*{3cm}\n"
        "\\noindent\\colorbox{docaccentdark}{\\parbox{\\dimexpr\\textwidth-2\\fboxsep\\relax}{\\centering\n"
        "\\vspace{1.5cm}\n"
        f"{{\\color{{white}}\\fontsize{{26}}{{32}}\\selectfont\\bfseries {_le(title)}}}\\\\[0.5cm]\n"
        "{\\color{docgold}\\rule{5cm}{1.4pt}}\\\\[0.4cm]\n"
        f"{client_line}"
        "\\vspace{1.3cm}\n"
        "}}\n"
        "\\vspace{1.4cm}\n"
        "\\renewcommand{\\arraystretch}{1.5}\n"
        "\\begin{tabular}{r l}\n"
        f"\\textcolor{{docmuted}}{{\\textbf{{{_le(labels['version'])}:}}}} & 1.0 \\\\\n"
        f"\\textcolor{{docmuted}}{{\\textbf{{{'Fecha' if document_language == 'es' else 'Date'}:}}}} & {timestamp} \\\\\n"
        "\\end{tabular}\n"
        "\\vfill\n"
        f"{ofi_logo_block}\n"
        "\\end{titlepage}\n\n"
        "\\clearpage\n"
        "\\tableofcontents\n"
        "\\clearpage\n\n"
        f"{body_block}\n"
        "\\end{document}\n"
    )


def _to_png_bytes(raw: bytes) -> bytes:
    """Normalise any PIL-readable image to PNG so reportlab and python-docx never choke."""
    from PIL import Image as PILImage
    buf = io.BytesIO(raw)
    img = PILImage.open(buf)
    img.load()
    if img.mode not in ("RGB", "RGBA", "L"):
        img = img.convert("RGBA")
    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


def _le(value: str) -> str:
    """Escape a string for LaTeX."""
    replacements = [
        ("\\", r"\textbackslash{}"),
        ("{", r"\{"), ("}", r"\}"), ("$", r"\$"), ("&", r"\&"),
        ("#", r"\#"), ("_", r"\_"), ("%", r"\%"),
        ("~", r"\textasciitilde{}"), ("^", r"\textasciicircum{}"),
    ]
    out = value
    for src, dst in replacements:
        out = out.replace(src, dst)
    return out


# ---------------------------------------------------------------------------
# PDF rendering
# ---------------------------------------------------------------------------

class _ReportDocTemplate(SimpleDocTemplate):
    """Feeds numbered section headings into the TableOfContents flowable and registers PDF
    outline bookmarks, so the compiled PDF gets a real, page-numbered TOC and a clickable
    sidebar outline. Requires doc.multiBuild() (two passes) instead of doc.build()."""

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph) and getattr(flowable, "style", None) is not None:
            if flowable.style.name == "SecH":
                text = flowable.getPlainText()
                self.notify("TOCEntry", (0, text, self.page))
                key = f"sec-{id(flowable)}-{self.page}"
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(text, key, level=0, closed=False)


def render_pdf(
    *,
    title: str,
    client_name: str,
    source_text: str,
    document_sections: list[DocSection],
    document_language: str,
    logo: Any,
    timestamp: str,
    image_data: list[dict[str, Any]] | None = None,
) -> bytes:
    image_data = image_data or []
    buffer = io.BytesIO()
    doc = _ReportDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1.8 * cm, rightMargin=1.8 * cm,
        topMargin=2.0 * cm, bottomMargin=1.7 * cm,
    )
    styles = getSampleStyleSheet()

    ink = colors.HexColor(_hex("ink"))
    ink_text = colors.HexColor(_hex("ink_text"))
    accent = colors.HexColor(_hex("accent"))
    accent_dark = colors.HexColor(_hex("accent_dark"))
    gold = colors.HexColor(_hex("gold"))
    muted = colors.HexColor(_hex("muted"))
    muted_light = colors.HexColor(_hex("muted_light"))
    border = colors.HexColor(_hex("border"))
    border_light = colors.HexColor(_hex("border_light"))
    band = colors.HexColor(_hex("band"))

    S = {
        "cover_title": ParagraphStyle("CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=27,
                                      leading=32, textColor=colors.white, alignment=TA_CENTER),
        "cover_sub": ParagraphStyle("CoverSub", parent=styles["BodyText"], fontName="Helvetica", fontSize=12.5,
                                    leading=17, textColor=colors.white, alignment=TA_CENTER),
        "info_label": ParagraphStyle("InfoLabel", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=9,
                                     leading=13, textColor=muted, alignment=TA_LEFT),
        "info_value": ParagraphStyle("InfoValue", parent=styles["BodyText"], fontName="Helvetica", fontSize=10,
                                     leading=13, textColor=ink_text, alignment=TA_LEFT),
        "toc_title": ParagraphStyle("TocTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=19,
                                    leading=23, textColor=ink, alignment=TA_LEFT, spaceAfter=2),
        "h1": ParagraphStyle("SecH", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=14.5,
                              leading=18, textColor=ink, spaceBefore=18, spaceAfter=6),
        "body": ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5,
                               leading=16, textColor=ink_text, spaceAfter=4),
        "bullet": ParagraphStyle("Bullet", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5,
                                 leading=16, textColor=ink_text, leftIndent=14, spaceAfter=3),
        "th": ParagraphStyle("TH", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=9.5,
                             textColor=colors.white),
        "td": ParagraphStyle("TD", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5,
                             textColor=ink_text),
        "caption": ParagraphStyle("FigCaption", parent=styles["BodyText"], fontName="Helvetica-Oblique",
                                  fontSize=9, leading=12, textColor=muted,
                                  alignment=TA_CENTER, spaceAfter=8),
    }

    story: list[Any] = []

    client_logo_bytes: io.BytesIO | None = None
    if logo:
        logo.seek(0)
        client_logo_bytes = io.BytesIO(logo.read())

    logo_path = resolve_ofi_logo_path()
    labels = localized_labels(document_language)
    ofi_logo_reader = ImageReader(str(logo_path)) if logo_path else None

    numbering = _numbered_titles(document_sections)
    available_width = A4[0] - doc.leftMargin - doc.rightMargin

    # ---- Cover page ----
    if client_logo_bytes:
        client_logo_bytes.seek(0)
        cover_logo = Image(client_logo_bytes, width=2.6 * cm, height=2.6 * cm, kind="proportional")
        cover_logo.hAlign = "CENTER"
        story.append(Spacer(1, 1.4 * cm))
        story.append(cover_logo)
        story.append(Spacer(1, 0.7 * cm))
    else:
        story.append(Spacer(1, 2.6 * cm))

    band_rows: list[Any] = [
        [Paragraph(_p(title or "Document"), S["cover_title"])],
        [HRFlowable(width="16%", thickness=1.6, color=gold, spaceBefore=10, spaceAfter=10, hAlign="CENTER")],
    ]
    if client_name:
        band_rows.append([Paragraph(_p(client_name), S["cover_sub"])])
    band_table = Table(band_rows, colWidths=[available_width])
    band_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), accent_dark),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, 0), 32),
        ("BOTTOMPADDING", (0, -1), (-1, -1), 32),
        ("LEFTPADDING", (0, 0), (-1, -1), 26), ("RIGHTPADDING", (0, 0), (-1, -1), 26),
    ]))
    story.append(band_table)
    story.append(Spacer(1, 1.2 * cm))

    date_label = "Fecha" if document_language == "es" else "Date"
    info_rows = [
        [Paragraph(labels["version"].upper(), S["info_label"]), Paragraph("1.0", S["info_value"])],
        [Paragraph(date_label.upper(), S["info_label"]), Paragraph(timestamp, S["info_value"])],
    ]
    info_table = Table(info_rows, colWidths=[available_width * 0.32, available_width * 0.68], hAlign="CENTER")
    info_table.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.8, border),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, border_light),
        ("LEFTPADDING", (0, 0), (-1, -1), 14), ("RIGHTPADDING", (0, 0), (-1, -1), 14),
        ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 2.4 * cm))
    if ofi_logo_reader is not None:
        cover_brand = Image(str(logo_path), width=1.9 * cm, height=1.9 * cm, kind="proportional")
        cover_brand.hAlign = "CENTER"
        story.append(cover_brand)
    story.append(PageBreak())

    # ---- Table of contents — own page, real page numbers via TableOfContents ----
    story.append(Paragraph(labels["contents"], S["toc_title"]))
    story.append(HRFlowable(width="100%", thickness=1.3, color=gold, spaceBefore=4, spaceAfter=16))
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("TOCLevel0", parent=styles["BodyText"], fontName="Helvetica", fontSize=11.5,
                       leading=22, textColor=ink_text, leftIndent=0, firstLineIndent=0),
    ]
    story.append(toc)
    story.append(PageBreak())

    # ---- Body sections ----
    for i, sec in enumerate(document_sections):
        sec_title = str(sec.get("title", ""))
        sec_type = str(sec.get("type", "paragraph"))
        content = sec.get("content")
        is_lead_summary = i == 0 and sec_type == "paragraph" and bool(sec_title)

        if sec_type == "figure" and isinstance(content, dict):
            fig_idx = int(content.get("figure_index", 0))
            caption = str(content.get("caption", ""))
            story.append(Spacer(1, 0.3 * cm))
            if fig_idx < len(image_data):
                try:
                    img_bytes = _to_png_bytes(image_data[fig_idx]["bytes"])
                    fig_img = Image(io.BytesIO(img_bytes), width=min(13 * cm, available_width), kind="proportional")
                    fig_img.hAlign = "CENTER"
                    story.append(fig_img)
                except Exception:
                    pass  # image unreadable; caption still renders below
            if caption:
                story.append(Paragraph(_p(caption), S["caption"]))
            story.append(Spacer(1, 0.4 * cm))
            continue

        if sec_title:
            story.append(Paragraph(_p(numbering.get(i, sec_title)), S["h1"]))
            story.append(HRFlowable(width="100%", thickness=1.1, color=gold, spaceBefore=0, spaceAfter=10))

        if is_lead_summary:
            summary_table = Table([[Paragraph(_p(str(content)), S["body"])]], colWidths=[available_width])
            summary_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), band),
                ("BOX", (0, 0), (-1, -1), 0.8, border),
                ("LEFTPADDING", (0, 0), (-1, -1), 14), ("RIGHTPADDING", (0, 0), (-1, -1), 14),
                ("TOPPADDING", (0, 0), (-1, -1), 12), ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
            ]))
            story.append(summary_table)

        elif sec_type == "bullets" and isinstance(content, list):
            for item in content:
                story.append(Paragraph(_p(str(item)), S["bullet"], bulletText="•"))

        elif sec_type == "table" and isinstance(content, dict):
            headers = [str(h) for h in content.get("headers", [])]
            rows = [[str(c) for c in row] for row in content.get("rows", [])]
            if headers:
                col_w = available_width / len(headers)
                tbl_data = [[Paragraph(_p(h), S["th"]) for h in headers]] + \
                           [[Paragraph(_p(c), S["td"]) for c in row] for row in rows]
                pdf_tbl = Table(tbl_data, colWidths=[col_w] * len(headers), repeatRows=1)
                pdf_tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), accent),
                    ("LINEBELOW", (0, 0), (-1, 0), 1.4, gold),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [band, colors.white]),
                    ("BOX", (0, 0), (-1, -1), 0.8, border),
                    ("INNERGRID", (0, 1), (-1, -1), 0.4, border_light),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                    ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]))
                story.append(pdf_tbl)

        else:
            story.append(Paragraph(_p(str(content or "")), S["body"]))

        story.append(Spacer(1, 0.25 * cm))

    def draw_cover(canvas, doc_obj) -> None:
        pass

    def draw_page(canvas, doc_obj) -> None:
        canvas.saveState()
        header_title = title if len(title) <= 70 else f"{title[:67]}..."
        canvas.setFont("Helvetica", 8.5)
        canvas.setFillColor(muted_light)
        canvas.drawString(doc_obj.leftMargin, A4[1] - 1.25 * cm, header_title)
        if ofi_logo_reader is not None:
            canvas.drawImage(
                ofi_logo_reader,
                A4[0] - doc_obj.rightMargin - 1.2 * cm, A4[1] - 1.55 * cm,
                width=1.2 * cm, height=1.2 * cm, preserveAspectRatio=True, mask="auto",
            )
        else:
            canvas.setFont("Helvetica-Bold", 9)
            canvas.setFillColor(accent)
            canvas.drawRightString(A4[0] - doc_obj.rightMargin, A4[1] - 1.3 * cm, "OFI")
        canvas.setStrokeColor(border)
        canvas.setLineWidth(0.8)
        canvas.line(doc_obj.leftMargin, A4[1] - 1.45 * cm, A4[0] - doc_obj.rightMargin, A4[1] - 1.45 * cm)
        canvas.setStrokeColor(border)
        canvas.line(doc_obj.leftMargin, 1.15 * cm, A4[0] - doc_obj.rightMargin, 1.15 * cm)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(muted_light)
        canvas.drawString(doc_obj.leftMargin, 0.85 * cm, "Documentator")
        canvas.drawRightString(A4[0] - doc_obj.rightMargin, 0.85 * cm, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    doc.multiBuild(story, onFirstPage=draw_cover, onLaterPages=draw_page)
    return buffer.getvalue()


def _p(value: str) -> str:
    return (
        value.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;")
             .replace("\n", "<br/>")
    )


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------

def _docx_shade(element_pr, hex_color: str) -> None:
    """Apply background shading to a cell (tcPr) or paragraph (pPr) properties element."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    element_pr.append(shd)


def _docx_cell_shade(cell, hex_color: str) -> None:
    _docx_shade(cell._tc.get_or_add_tcPr(), hex_color)


def _docx_paragraph_border(paragraph, *, color: str, size: int = 10) -> None:
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pBdr)


def _docx_add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    """Insert a real Word field (TOC, PAGE, ...) that Word computes/updates on open."""
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(begin)
    r.append(instr)
    r.append(separate)
    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        r.append(t)
    r.append(end)


def _docx_set_run_color(run, hex_color: str) -> None:
    run.font.color.rgb = RGBColor.from_string(hex_color)


def _docx_configure_styles(document: Document) -> None:
    """Base typography: Calibri body copy, teal Heading 1 with a gold rule, generous line spacing."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(PALETTE["ink_text"])
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(8)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor.from_string(PALETTE["ink"])
    heading1.paragraph_format.space_before = Pt(20)
    heading1.paragraph_format.space_after = Pt(4)


def render_docx(
    *,
    title: str,
    client_name: str,
    document_sections: list[DocSection],
    document_language: str,
    timestamp: str,
    image_data: list[dict[str, Any]] | None = None,
) -> bytes:
    image_data = image_data or []
    document = Document()
    labels = localized_labels(document_language)
    logo_path = resolve_ofi_logo_path()
    numbering = _numbered_titles(document_sections)
    _docx_configure_styles(document)

    section = document.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    # ---- Running header: document title (left) + OFI brand (right) ----
    header_title = title if len(title) <= 70 else f"{title[:67]}..."
    hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    hp.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
    title_run = hp.add_run(header_title)
    title_run.font.size = Pt(8.5)
    _docx_set_run_color(title_run, PALETTE["muted_light"])
    hp.add_run("\t")
    if logo_path:
        hp.add_run().add_picture(str(logo_path), height=Cm(0.9))
    else:
        brand_run = hp.add_run("OFI")
        brand_run.bold = True
        _docx_set_run_color(brand_run, PALETTE["accent"])
    _docx_paragraph_border(hp, color=PALETTE["border"], size=6)

    # ---- Running footer: brand (left) + page number (right) ----
    fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    fp.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
    foot_run = fp.add_run("Documentator")
    foot_run.font.size = Pt(8)
    _docx_set_run_color(foot_run, PALETTE["muted_light"])
    fp.add_run("\t")
    page_label_run = fp.add_run(f"{'Página' if document_language == 'es' else 'Page'} ")
    page_label_run.font.size = Pt(8)
    _docx_set_run_color(page_label_run, PALETTE["muted_light"])
    _docx_add_field(fp, "PAGE", "1")

    # ---- Cover page ----
    if logo_path:
        logo_para = document.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.add_run().add_picture(str(logo_path), height=Cm(1.6))
        logo_para.paragraph_format.space_after = Pt(18)

    band_table = document.add_table(rows=2 if client_name else 1, cols=1)
    band_table.autofit = False
    band_table.columns[0].width = content_width
    title_cell = band_table.rows[0].cells[0]
    title_cell.width = content_width
    _docx_cell_shade(title_cell, PALETTE["accent_dark"])
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(28)
    title_para.paragraph_format.space_after = Pt(6) if client_name else Pt(28)
    title_run = title_para.add_run(title or "Document")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    _docx_set_run_color(title_run, PALETTE["white"])
    if client_name:
        sub_cell = band_table.rows[1].cells[0]
        _docx_cell_shade(sub_cell, PALETTE["accent_dark"])
        sub_para = sub_cell.paragraphs[0]
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.paragraph_format.space_after = Pt(28)
        sub_run = sub_para.add_run(client_name)
        sub_run.font.size = Pt(13)
        _docx_set_run_color(sub_run, PALETTE["white"])

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(20)

    info_table = document.add_table(rows=2, cols=2)
    info_table.autofit = False
    info_table.style = "Table Grid"
    info_labels = [labels["version"].upper(), ("FECHA" if document_language == "es" else "DATE")]
    info_values = ["1.0", timestamp]
    for row_idx, (label_text, value_text) in enumerate(zip(info_labels, info_values)):
        label_cell, value_cell = info_table.rows[row_idx].cells
        label_cell.width = Cm(3.5)
        value_cell.width = content_width - Cm(3.5)
        label_run = label_cell.paragraphs[0].add_run(label_text)
        label_run.font.bold = True
        label_run.font.size = Pt(9)
        _docx_set_run_color(label_run, PALETTE["muted"])
        value_run = value_cell.paragraphs[0].add_run(value_text)
        value_run.font.size = Pt(10)
    document.add_page_break()

    # ---- Table of contents — real Word TOC field ----
    toc_heading = document.add_paragraph()
    toc_heading.paragraph_format.space_after = Pt(4)
    toc_run = toc_heading.add_run(labels["contents"])
    toc_run.font.size = Pt(19)
    toc_run.font.bold = True
    _docx_set_run_color(toc_run, PALETTE["ink"])
    _docx_paragraph_border(toc_heading, color=PALETTE["gold"], size=16)

    toc_paragraph = document.add_paragraph()
    hint = "Haz clic derecho y elige “Actualizar campos” para generar el índice." if document_language == "es" \
        else "Right-click and choose “Update Field” to generate the table of contents."
    _docx_add_field(toc_paragraph, 'TOC \\o "1-1" \\h \\z \\u', hint)
    document.add_page_break()

    # ---- Body sections ----
    for i, sec in enumerate(document_sections):
        sec_title = str(sec.get("title", ""))
        sec_type = str(sec.get("type", "paragraph"))
        content = sec.get("content")
        is_lead_summary = i == 0 and sec_type == "paragraph" and bool(sec_title)

        if sec_type == "figure" and isinstance(content, dict):
            fig_idx = int(content.get("figure_index", 0))
            caption = str(content.get("caption", ""))
            if fig_idx < len(image_data):
                try:
                    img_bytes = _to_png_bytes(image_data[fig_idx]["bytes"])
                    pic_para = document.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic_para.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(14))
                except Exception:
                    pass  # image unreadable; caption still renders below
            if caption:
                cap_para = document.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(caption)
                cap_run.font.size = Pt(9)
                cap_run.font.italic = True
                _docx_set_run_color(cap_run, PALETTE["muted"])
            continue

        if sec_title:
            document.add_heading(numbering.get(i, sec_title), level=1)

        if is_lead_summary:
            box = document.add_table(rows=1, cols=1)
            box.autofit = False
            box.columns[0].width = content_width
            cell = box.rows[0].cells[0]
            cell.width = content_width
            _docx_cell_shade(cell, PALETTE["band"])
            cell_para = cell.paragraphs[0]
            cell_para.paragraph_format.space_before = Pt(4)
            cell_para.paragraph_format.space_after = Pt(4)
            cell_para.add_run(str(content))

        elif sec_type == "bullets" and isinstance(content, list):
            for item in content:
                document.add_paragraph(str(item), style="List Bullet")

        elif sec_type == "table" and isinstance(content, dict):
            headers = [str(h) for h in content.get("headers", [])]
            rows = [[str(c) for c in row] for row in content.get("rows", [])]
            if headers:
                tbl = document.add_table(rows=1 + len(rows), cols=len(headers))
                tbl.style = "Table Grid"
                tbl.autofit = True
                for ci, h in enumerate(headers):
                    header_cell = tbl.rows[0].cells[ci]
                    _docx_cell_shade(header_cell, PALETTE["accent"])
                    run = header_cell.paragraphs[0].add_run(h)
                    run.font.bold = True
                    _docx_set_run_color(run, PALETTE["white"])
                for ri, row_data in enumerate(rows):
                    band_this_row = ri % 2 == 0
                    for ci, cell_val in enumerate(row_data):
                        data_cell = tbl.rows[ri + 1].cells[ci]
                        if band_this_row:
                            _docx_cell_shade(data_cell, PALETTE["band"])
                        data_cell.paragraphs[0].add_run(cell_val)

        else:
            document.add_paragraph(str(content or ""))

    # Tell Word to recompute all fields (TOC, PAGE) as soon as the document is opened, instead
    # of showing raw field placeholders until the user manually chooses "Update Field".
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    document.settings.element.append(update_fields)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def _env_value(name: str, default: str = "") -> str:
    return os.getenv(name, default) or default


def _fallback_title(source_text: str) -> str:
    first = next((ln.strip(" :-") for ln in source_text.splitlines() if ln.strip()), "Document")
    return first[:80] or "Document"


def _fallback_client_name(source_text: str) -> str:
    match = re.search(r"(?:client|customer|company|for)[:\-]\s*([^\n]+)", source_text, re.IGNORECASE)
    return match.group(1).strip() if match else ""


def _fallback_summary(source_text: str, sentence_count: int = 2) -> str:
    parts = re.split(r"(?<=[.!?])\s+", re.sub(r"\s+", " ", source_text.strip()))
    return " ".join(parts[:sentence_count]) if parts else source_text[:240].strip() or "Generated documentation."


def _extract_lines(source_text: str) -> list[str]:
    lines = []
    for raw in source_text.splitlines():
        clean = re.sub(r"^[-*•\d+.\)]\s*", "", raw.strip())
        if clean:
            lines.append(clean)
    return lines


def build_tex_zip(
    latex_source: str,
    filename: str,
    image_data: list[dict[str, Any]],
) -> bytes:
    """Bundle LaTeX source + all referenced images into a self-contained zip."""
    tex_name = f"{filename}.tex"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(tex_name, latex_source.encode("utf-8"))

        for i, img in enumerate(image_data):
            try:
                zf.writestr(f"figure_{i}.png", _to_png_bytes(img["bytes"]))
            except Exception:
                pass

        logo_path = resolve_ofi_logo_path()
        if logo_path and logo_path.exists():
            try:
                with open(logo_path, "rb") as f:
                    zf.writestr("ofi-logo.png", _to_png_bytes(f.read()))
            except Exception:
                pass

        zf.writestr("compile.sh", f"#!/bin/bash\npdflatex {tex_name}\npdflatex {tex_name}\n")
        zf.writestr("compile.bat", f"@echo off\npdflatex {tex_name}\npdflatex {tex_name}\n")

    return buf.getvalue()


def pdf_base64(pdf_bytes: bytes) -> str:
    return base64.b64encode(pdf_bytes).decode("ascii")


def docx_base64(docx_bytes: bytes) -> str:
    return base64.b64encode(docx_bytes).decode("ascii")


def tex_zip_base64(tex_zip_bytes: bytes) -> str:
    return base64.b64encode(tex_zip_bytes).decode("ascii")


def detect_document_language(source_text: str, agent_instructions: str = "") -> str:
    combined = f"{source_text}\n{agent_instructions}".lower()
    spanish_hits = sum(
        word in combined
        for word in [" el ", " la ", " los ", " las ", " para ", " documento ", " cliente ", " seccion ", " secciones ", " que ", " con "]
    )
    english_hits = sum(
        word in combined
        for word in [" the ", " document ", " client ", " section ", " sections ", " with ", " for ", " and ", " requirements "]
    )
    return "es" if spanish_hits >= english_hits else "en"


def localized_labels(language: str) -> dict[str, Any]:
    if language == "es":
        return {
            "contents": "Contenido",
            "client": "Cliente",
            "version": "Versión",
            "executive_summary": "Resumen Ejecutivo",
            "key_points": "Puntos Clave",
            "version_control": "Control de Versiones",
            "version_headers": ["Versión", "Fecha", "Autor", "Notas"],
            "initial_draft": "Borrador inicial",
        }
    return {
        "contents": "Contents",
        "client": "Client",
        "version": "Version",
        "executive_summary": "Executive Summary",
        "key_points": "Key Points",
        "version_control": "Version Control",
        "version_headers": ["Version", "Date", "Author", "Notes"],
        "initial_draft": "Initial draft",
    }


def resolve_ofi_logo_path() -> Path | None:
    candidates = [
        ASSETS_DIR / "ofi-logo.png",
        ASSETS_DIR / "ofi-logo.png.png",
        ASSETS_DIR / "ofi-logo.jpg",
        ASSETS_DIR / "ofi-logo.jpeg",
        ASSETS_DIR / "ofi-logo.webp",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate

    prefix_matches = sorted(ASSETS_DIR.glob("ofi-logo*")) if ASSETS_DIR.exists() else []
    return prefix_matches[0] if prefix_matches else None


def extract_pdf_texts(pdf_files: list[Any]) -> list[str]:
    """Extract plain text from a list of uploaded PDF file objects."""
    texts: list[str] = []
    for pdf_file in pdf_files:
        try:
            pdf_file.seek(0)
            reader = PdfReader(io.BytesIO(pdf_file.read()))
            pages = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(page_text.strip())
            if pages:
                texts.append("\n\n".join(pages))
        except Exception:
            pass
    return texts
